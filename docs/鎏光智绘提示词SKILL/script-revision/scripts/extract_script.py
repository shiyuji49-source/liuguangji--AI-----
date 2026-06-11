#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract_script.py —— 微短剧多集剧本摄入与分集切分

用法:
    python extract_script.py <文件或文件夹> [--out OUTDIR]

支持 .docx / .txt / .md / .pdf。自动识别 "第X集 / 第N集 / 【第一集】 / EP1" 等
分集标记切分；若一个文件夹内一集一个文件、且文件内无标记，则按文件切分并尝试
从文件名取集号。输出 OUTDIR/episodes.json 并打印总览。
"""
import argparse, json, os, re, sys
from pathlib import Path

SUPPORTED = {'.docx', '.txt', '.md', '.pdf'}

# ---------- 文本抽取 ----------
def read_txt(p: Path) -> str:
    for enc in ('utf-8', 'utf-8-sig', 'gb18030', 'gbk'):
        try:
            return p.read_text(encoding=enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return p.read_bytes().decode('utf-8', errors='replace')

def read_docx(p: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        sys.exit("缺少 python-docx，请先安装: pip install python-docx --break-system-packages")
    d = docx.Document(str(p))
    lines = [para.text for para in d.paragraphs]
    # 表格里也可能有内容
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    return "\n".join(lines)

def read_pdf(p: Path) -> str:
    try:
        import pdfplumber
        out = []
        with pdfplumber.open(str(p)) as pdf:
            for page in pdf.pages:
                out.append(page.extract_text() or "")
        return "\n".join(out)
    except ImportError:
        pass
    try:
        from pypdf import PdfReader
        r = PdfReader(str(p))
        return "\n".join((pg.extract_text() or "") for pg in r.pages)
    except ImportError:
        sys.exit("缺少 pdfplumber/pypdf，请先安装: pip install pdfplumber --break-system-packages")

def extract_text(p: Path) -> str:
    ext = p.suffix.lower()
    if ext == '.docx':
        return read_docx(p)
    if ext == '.pdf':
        return read_pdf(p)
    return read_txt(p)

# ---------- 中文数字 ----------
_CN_DIGIT = {'零':0,'〇':0,'一':1,'二':2,'两':2,'三':3,'四':4,'五':5,
             '六':6,'七':7,'八':8,'九':9}
_CN_UNIT = {'十':10,'百':100,'千':1000}

def cn2int(s: str):
    s = s.strip()
    if s.isdigit():
        return int(s)
    total, num = 0, 0
    for ch in s:
        if ch in _CN_DIGIT:
            num = _CN_DIGIT[ch]
        elif ch in _CN_UNIT:
            unit = _CN_UNIT[ch]
            total += (num if num else 1) * unit
            num = 0
    total += num
    return total or None

# ---------- 分集切分 ----------
EP_RE = re.compile(
    r'^[\s\u3010\uFF3B\(\[【〔]*第\s*([0-9零〇一二两三四五六七八九十百千]+)\s*集'
    r'|^[\s\u3010\uFF3B\(\[【〔]*(?:EP|Ep|ep|E)\s*[\.\-:：]?\s*([0-9]+)\b'
)
NUM_IN_NAME = re.compile(r'(\d+)')

# 场景头估计: 含 内/外 且含 日/夜 时段, 或以 "场N" 开头
SCENE_HEAD_RE = re.compile(r'(^|\s)场\s*\d+|(内景|外景|内|外).{0,12}(日|夜|晨|黄昏|傍晚|清晨|凌晨|白天|深夜)')
NIGHT_RE = re.compile(r'(夜|晚|深夜|凌晨|傍晚)')
DAY_RE = re.compile(r'(日|白天|清晨|晨|黄昏)')

def split_episodes_in_text(text, source):
    """在单个文件的文本里按标记切分。返回 [(ep_number_or_None, label, title, body)]"""
    lines = text.splitlines()
    marks = []  # (line_idx, number, label, title)
    for i, ln in enumerate(lines):
        m = EP_RE.search(ln)
        if m:
            raw = m.group(1) or m.group(2)
            num = cn2int(raw)
            label = ln.strip()
            title = ln[m.end():].strip(' 　:：·-—【】［］')
            marks.append((i, num, label, title))
    if not marks:
        return [(None, None, "", text)]
    eps = []
    for j, (idx, num, label, title) in enumerate(marks):
        end = marks[j+1][0] if j+1 < len(marks) else len(lines)
        body = "\n".join(lines[idx+1:end]).strip()
        eps.append((num, label, title, body))
    return eps

def analyze_body(body):
    wc = len(re.sub(r'\s', '', body))
    scenes = day = night = 0
    for ln in body.splitlines():
        if SCENE_HEAD_RE.search(ln):
            scenes += 1
            if NIGHT_RE.search(ln):
                night += 1
            elif DAY_RE.search(ln):
                day += 1
    return wc, scenes, day, night

def natural_key(p: Path):
    nums = NUM_IN_NAME.findall(p.stem)
    return (int(nums[0]) if nums else 1_000_000, p.name)

# ---------- 主流程 ----------
def gather_files(target: Path):
    if target.is_dir():
        files = sorted([p for p in target.iterdir()
                        if p.suffix.lower() in SUPPORTED], key=natural_key)
        if not files:
            sys.exit(f"文件夹内没有支持的剧本文件: {target}")
        return files
    if target.suffix.lower() not in SUPPORTED:
        sys.exit(f"不支持的格式: {target.suffix}（支持 {', '.join(sorted(SUPPORTED))}）")
    return [target]

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('target', help='剧本文件或文件夹')
    ap.add_argument('--out', default='/tmp/script_parsed', help='输出目录')
    args = ap.parse_args()

    target = Path(args.target).expanduser()
    if not target.exists():
        sys.exit(f"路径不存在: {target}")
    outdir = Path(args.out).expanduser()
    outdir.mkdir(parents=True, exist_ok=True)

    files = gather_files(target)
    warnings = []
    raw_eps = []  # (number_or_None, label, title, body, source)
    for f in files:
        text = extract_text(f)
        if not text.strip():
            warnings.append(f"{f.name} 抽取为空（可能是扫描版 PDF，需 OCR）")
            continue
        eps = split_episodes_in_text(text, f.name)
        if eps[0][0] is None and len(eps) == 1:
            # 文件内无分集标记：整文件当一集，尝试从文件名取号
            nm = NUM_IN_NAME.findall(f.stem)
            num = int(nm[0]) if nm else None
            raw_eps.append((num, f.stem, "", eps[0][3], f.name))
        else:
            for (num, label, title, body) in eps:
                raw_eps.append((num, label, title, body, f.name))

    if not raw_eps:
        sys.exit("没有解析出任何内容。")

    # 排序：有集号按集号，否则保持出现顺序
    if all(e[0] is not None for e in raw_eps):
        raw_eps.sort(key=lambda e: e[0])
    has_markers = any(e[0] is not None for e in raw_eps)
    if not has_markers:
        warnings.append("未发现任何分集标记（第X集 / EP1 等）。已按文件/整体顺序处理；"
                        "若这是小说或大纲，请先与用户确认是否需要‘小说转剧本’。")

    episodes = []
    total_words = 0
    for i, (num, label, title, body, source) in enumerate(raw_eps, 1):
        wc, scenes, day, night = analyze_body(body)
        total_words += wc
        episodes.append({
            "index": i,
            "ep_number": num if num is not None else i,
            "ep_label": label or f"第{i}集",
            "title": title,
            "source_file": source,
            "word_count": wc,
            "scene_count_est": scenes,
            "day_scenes_est": day,
            "night_scenes_est": night,
            "text": body,
        })

    meta = {
        "total_episodes": len(episodes),
        "total_words": total_words,
        "file_count": len(files),
        "has_episode_markers": has_markers,
        "warnings": warnings,
    }
    (outdir / "episodes.json").write_text(
        json.dumps({"meta": meta, "episodes": episodes}, ensure_ascii=False, indent=2),
        encoding='utf-8')

    # ---------- 总览 ----------
    print("=" * 56)
    print(f"摄入完成 → {outdir/'episodes.json'}")
    print(f"文件数: {meta['file_count']}   总集数: {meta['total_episodes']}   "
          f"总字数: {meta['total_words']:,}")
    avg = total_words // max(len(episodes), 1)
    print(f"单集均字数: 约 {avg}")
    if total_words > 120000:
        print("⚠ 体量较大（>12万字），建议分批通读、增量更新剧本圣经。")
    if warnings:
        print("\n提示:")
        for w in warnings:
            print(f"  - {w}")
    print("\n前几集概览（集号 | 字数 | 场次估 | 标题）:")
    for ep in episodes[:8]:
        print(f"  第{ep['ep_number']}集 | {ep['word_count']:>5}字 | "
              f"{ep['scene_count_est']}场 | {ep['title'][:20]}")
    if len(episodes) > 8:
        print(f"  …… 共 {len(episodes)} 集")
    print("=" * 56)

if __name__ == '__main__':
    main()
